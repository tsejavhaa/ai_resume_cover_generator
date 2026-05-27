"""
Export service — Phase 3.

Converts a GenerateResponse into downloadable files:
  - PDF  → cover letter + resume tweaks via reportlab
  - DOCX → cover letter + resume tweaks via python-docx
  - MD   → plain markdown (no dependencies)
"""
import io
from loguru import logger
from app.models.schemas import GenerateResponse


# ── PDF Export ────────────────────────────────────────────────

def export_pdf(result: GenerateResponse) -> bytes:
    """Render cover letter and resume tweaks as a styled PDF."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer,
        HRFlowable, Table, TableStyle,
    )

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
    )

    styles = getSampleStyleSheet()
    DARK = colors.HexColor("#1a1a2e")
    ACCENT = colors.HexColor("#4f46e5")
    MUTED = colors.HexColor("#6b7280")

    title_style = ParagraphStyle(
        "Title", parent=styles["Heading1"],
        fontSize=20, textColor=DARK, spaceAfter=4,
    )
    subtitle_style = ParagraphStyle(
        "Subtitle", parent=styles["Normal"],
        fontSize=10, textColor=MUTED, spaceAfter=16,
    )
    section_style = ParagraphStyle(
        "Section", parent=styles["Heading2"],
        fontSize=13, textColor=ACCENT, spaceBefore=18, spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontSize=10, leading=16, textColor=DARK,
    )
    label_style = ParagraphStyle(
        "Label", parent=styles["Normal"],
        fontSize=9, textColor=MUTED, spaceBefore=4,
    )
    tweak_orig_style = ParagraphStyle(
        "TweakOrig", parent=styles["Normal"],
        fontSize=9, textColor=colors.HexColor("#dc2626"),
        leading=14,
    )
    tweak_new_style = ParagraphStyle(
        "TweakNew", parent=styles["Normal"],
        fontSize=9, textColor=colors.HexColor("#16a34a"),
        leading=14,
    )

    story = []

    # ── Header ────────────────────────────────────────────────
    story.append(Paragraph("Cover Letter & Resume Report", title_style))
    score_color = "#16a34a" if result.match_score >= 60 else "#d97706" if result.match_score >= 40 else "#dc2626"
    story.append(Paragraph(
        f'Match Score: <font color="{score_color}"><b>{result.match_score}%</b></font> '
        f'&nbsp;|&nbsp; Model: {result.model_used} &nbsp;|&nbsp; Backend: {result.backend_used}',
        subtitle_style,
    ))
    story.append(HRFlowable(width="100%", thickness=2, color=ACCENT, spaceAfter=16))

    # ── Skills summary ────────────────────────────────────────
    story.append(Paragraph("Skills Analysis", section_style))
    skills_data = [
        ["Matched Skills", "Missing Skills"],
        [
            ", ".join(result.matched_skills) or "None",
            ", ".join(result.missing_skills) or "None",
        ],
    ]
    skills_table = Table(skills_data, colWidths=[3 * inch, 3 * inch])
    skills_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), ACCENT),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, 0), 9),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 1), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f8fafc"), colors.white]),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#e2e8f0")),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(skills_table)

    # ── Cover Letter ──────────────────────────────────────────
    story.append(Paragraph("Cover Letter", section_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e2e8f0"), spaceAfter=10))
    for para in result.cover_letter.split("\n\n"):
        if para.strip():
            story.append(Paragraph(para.strip().replace("\n", " "), body_style))
            story.append(Spacer(1, 8))

    # ── Resume Tweaks ─────────────────────────────────────────
    if result.resume_tweaks:
        story.append(Paragraph("Resume Tweaks", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e2e8f0"), spaceAfter=10))
        for i, tweak in enumerate(result.resume_tweaks, 1):
            story.append(Paragraph(f"{i}. [{tweak.section}] — {tweak.reason}", label_style))
            if tweak.original:
                story.append(Paragraph(f"Before: {tweak.original}", tweak_orig_style))
            if tweak.suggested:
                story.append(Paragraph(f"After:  {tweak.suggested}", tweak_new_style))
            story.append(Spacer(1, 10))

    doc.build(story)
    return buffer.getvalue()


# ── DOCX Export ───────────────────────────────────────────────

def export_docx(result: GenerateResponse) -> bytes:
    """Render cover letter and resume tweaks as a DOCX document."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    def add_heading(text: str, level: int = 1, color: tuple = (79, 70, 229)):
        h = doc.add_heading(text, level=level)
        h.runs[0].font.color.rgb = RGBColor(*color)
        return h

    def add_body(text: str):
        p = doc.add_paragraph(text)
        p.runs[0].font.size = Pt(10.5)
        return p

    # Title
    title = doc.add_heading("Cover Letter & Resume Report", 0)
    title.runs[0].font.color.rgb = RGBColor(26, 26, 46)

    # Match score
    score_para = doc.add_paragraph()
    score_run = score_para.add_run(f"Match Score: {result.match_score}%")
    score_run.bold = True
    score_run.font.size = Pt(11)
    meta_run = score_para.add_run(
        f"  |  Model: {result.model_used}  |  Backend: {result.backend_used}"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_paragraph()

    # Skills
    add_heading("Skills Analysis", level=1)
    matched = doc.add_paragraph()
    matched.add_run("Matched: ").bold = True
    matched.add_run(", ".join(result.matched_skills) or "None")
    missing = doc.add_paragraph()
    missing.add_run("Missing: ").bold = True
    run = missing.add_run(", ".join(result.missing_skills) or "None")
    if result.missing_skills:
        run.font.color.rgb = RGBColor(220, 38, 38)

    doc.add_paragraph()

    # Cover letter
    add_heading("Cover Letter", level=1)
    for para in result.cover_letter.split("\n\n"):
        if para.strip():
            add_body(para.strip().replace("\n", " "))

    doc.add_paragraph()

    # Resume tweaks
    if result.resume_tweaks:
        add_heading("Resume Tweaks", level=1)
        for i, tweak in enumerate(result.resume_tweaks, 1):
            section_p = doc.add_paragraph()
            section_p.add_run(f"{i}. [{tweak.section}]").bold = True
            section_p.add_run(f" — {tweak.reason}")
            section_p.runs[-1].font.size = Pt(10)

            if tweak.original:
                before = doc.add_paragraph()
                before.add_run("Before: ").bold = True
                r = before.add_run(tweak.original)
                r.font.color.rgb = RGBColor(220, 38, 38)
                r.font.size = Pt(9.5)

            if tweak.suggested:
                after = doc.add_paragraph()
                after.add_run("After: ").bold = True
                r = after.add_run(tweak.suggested)
                r.font.color.rgb = RGBColor(22, 163, 74)
                r.font.size = Pt(9.5)

            doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ── Markdown Export ───────────────────────────────────────────

def export_markdown(result: GenerateResponse) -> bytes:
    lines = [
        "# Cover Letter & Resume Report",
        "",
        f"**Match Score:** {result.match_score}% | "
        f"**Model:** {result.model_used} | **Backend:** {result.backend_used}",
        "",
        f"**Matched Skills:** {', '.join(result.matched_skills) or 'None'}",
        f"**Missing Skills:** {', '.join(result.missing_skills) or 'None'}",
        "",
        "---",
        "",
        "## Cover Letter",
        "",
        result.cover_letter.strip(),
        "",
        "---",
        "",
        "## Resume Tweaks",
        "",
    ]
    for i, tweak in enumerate(result.resume_tweaks, 1):
        lines += [
            f"### {i}. [{tweak.section}]",
            f"**Reason:** {tweak.reason}",
            "",
            f"**Before:** {tweak.original}",
            "",
            f"**After:** {tweak.suggested}",
            "",
        ]
    return "\n".join(lines).encode("utf-8")