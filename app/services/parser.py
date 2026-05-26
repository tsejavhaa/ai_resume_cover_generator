"""
Document parser: extract raw text from uploaded resume files.
Supports PDF, DOCX, and plain text.
"""
import io
from pathlib import Path
from loguru import logger


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Dispatch to the correct parser based on file extension.
    Returns cleaned plain text.
    """
    ext = Path(filename).suffix.lower()
    logger.info(f"Parsing document: {filename} ({len(file_bytes)} bytes, type={ext})")

    if ext == ".pdf":
        return _parse_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return _parse_docx(file_bytes)
    elif ext in (".txt", ".md"):
        return _parse_text(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            "Please upload a PDF, DOCX, or plain text file."
        )


def _parse_pdf(file_bytes: bytes) -> str:
    from pdfminer.high_level import extract_text_to_fp
    from pdfminer.layout import LAParams

    output = io.StringIO()
    extract_text_to_fp(
        io.BytesIO(file_bytes),
        output,
        laparams=LAParams(),
        output_type="text",
        codec=None,
    )
    text = output.getvalue()
    return _clean(text)


def _parse_docx(file_bytes: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also grab text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return _clean("\n".join(paragraphs))


def _parse_text(file_bytes: bytes) -> str:
    try:
        return _clean(file_bytes.decode("utf-8"))
    except UnicodeDecodeError:
        return _clean(file_bytes.decode("latin-1"))


def _clean(text: str) -> str:
    """Remove excessive whitespace while preserving paragraph structure."""
    lines = [line.strip() for line in text.splitlines()]
    # Collapse 3+ consecutive blank lines into 2
    cleaned, blank_count = [], 0
    for line in lines:
        if line == "":
            blank_count += 1
            if blank_count <= 2:
                cleaned.append(line)
        else:
            blank_count = 0
            cleaned.append(line)
    return "\n".join(cleaned).strip()